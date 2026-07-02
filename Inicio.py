import streamlit as st
from utils import aplicar_estilos
from database import init_db, get_equipos, get_calibraciones, get_recordatorios_calibracion
from datetime import datetime, date
import io
from docx import Document
from docx.shared import RGBColor, Pt


def generar_reporte_recordatorios_docx(tabla_recordatorios):
    buffer = io.BytesIO()
    doc = Document()

    titulo = doc.add_heading("Reporte de recordatorios de calibración", level=1)
    titulo.runs[0].font.color.rgb = RGBColor(0x0b, 0x2c, 0x40)

    p = doc.add_paragraph()
    p.add_run("Generado por MetriCore").bold = True
    p.add_run(f" · Fecha de generación: {date.today().strftime('%d/%m/%Y')}")

    doc.add_paragraph(
        "Este reporte lista los equipos que requieren calibración o que tienen una "
        "calibración próxima dentro de los siguientes 30 días."
    )

    if tabla_recordatorios.empty:
        doc.add_paragraph("No hay recordatorios pendientes.")
    else:
        table = doc.add_table(rows=1, cols=len(tabla_recordatorios.columns))
        table.style = "Table Grid"
        for i, col in enumerate(tabla_recordatorios.columns):
            cell = table.rows[0].cells[i]
            cell.text = str(col)
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(9)

        for _, row in tabla_recordatorios.iterrows():
            cells = table.add_row().cells
            for i, col in enumerate(tabla_recordatorios.columns):
                cells[i].text = "" if row[col] is None else str(row[col])
                for run in cells[i].paragraphs[0].runs:
                    run.font.size = Pt(8)

    doc.save(buffer)
    return buffer.getvalue()

st.set_page_config(
    page_title="MetriCore",
    page_icon="🔬",
    layout="wide"
)
aplicar_estilos()

st.markdown("""
<style>
[data-testid="stSidebar"] {
    background-color: #0b2c40 !important;
}
[data-testid="stSidebar"] * {
    color: #E0F0ED !important;
}
[data-testid="stSidebarNav"] a {
    border-radius: 8px;
    padding: 6px 12px;
    margin: 2px 0;
}
[data-testid="stSidebarNav"] a:hover {
    background-color: #0a453c !important;
}
[data-testid="stSidebarNav"] a[aria-selected="true"] {
    background-color: #238d93 !important;
}
[data-testid="stMetric"] {
    background-color: #FFFFFF;
    border: 1px solid #2dc197;
    border-left: 5px solid #1469aa;
    border-radius: 8px;
    padding: 16px 20px;
}
[data-testid="stMetricLabel"] {
    color: #0a453c !important;
    font-size: 20px !important;
    font-weight: 500 !important;
}
[data-testid="stMetricValue"] {
    color: #0b2c40 !important;
    font-size: 36px !important;
    font-weight: 700 !important;
}
</style>
""", unsafe_allow_html=True)

init_db()

st.markdown("""
<div style="background: linear-gradient(135deg, #063d7d 0%, #238d93 60%, #15924a 100%);
border-radius: 12px; padding: 40px 48px; margin-bottom: 24px;">
    <h1 style="color:#FFFFFF; margin:0; font-size:36px;">🔬 MetriCore</h1>
    <p style="color:#E0F0ED; margin:8px 0 0 0; font-size:16px;">
    Plataforma integral de gestión y control metrológico — ISO/IEC 17025</p>
</div>
""", unsafe_allow_html=True)

equipos = get_equipos()
calibraciones = get_calibraciones()

total_equipos = len(equipos)
operativos = len(equipos[equipos["estado"] == "En servicio"]) if not equipos.empty else 0
porcentaje_op = round((operativos / total_equipos * 100)) if total_equipos > 0 else 0

calibraciones_proximas = 0
if not equipos.empty and "proxima_calibracion" in equipos.columns:
    hoy = date.today()
    for _, row in equipos.iterrows():
        if row["proxima_calibracion"]:
            try:
                fecha = datetime.strptime(row["proxima_calibracion"], "%Y-%m-%d").date()
                dias = (fecha - hoy).days
                if 0 <= dias <= 30:
                    calibraciones_proximas += 1
            except:
                pass

col1, col2, col3, col4 = st.columns(4)
with col1:
    st.metric("🔧 Total equipos", total_equipos, help="Instrumentos registrados")
with col2:
    st.metric("✅ Equipos operativos", f"{porcentaje_op}%")
with col3:
    st.metric("📅 Calibraciones próximas", calibraciones_proximas, help="Vencen en 30 días")
with col4:
    st.metric("📋 Total calibraciones", len(calibraciones))

st.divider()

recordatorios = get_recordatorios_calibracion(dias_alerta=30)

st.markdown("<h3 style='color:#0b2c40; font-size:24px;'>Recordatorios de calibración</h3>", unsafe_allow_html=True)
if recordatorios.empty:
    st.success("No hay equipos que requieran calibración ni calibraciones próximas en los siguientes 30 días.")
else:
    requieren_calibracion = len(recordatorios[recordatorios["dias_restantes"] < 0])
    hoy_vencen = len(recordatorios[recordatorios["dias_restantes"] == 0])
    semana = len(recordatorios[(recordatorios["dias_restantes"] > 0) & (recordatorios["dias_restantes"] <= 7)])
    mes = len(recordatorios[(recordatorios["dias_restantes"] > 7) & (recordatorios["dias_restantes"] <= 30)])

    r1, r2, r3, r4 = st.columns(4)
    r1.metric("Requieren calibración", requieren_calibracion)
    r2.metric("Vencen hoy", hoy_vencen)
    r3.metric("Próximos 7 días", semana)
    r4.metric("Próximos 30 días", mes)

    if requieren_calibracion or hoy_vencen:
        st.error("Hay equipos que requieren calibración o que vencen hoy. Revise la tabla de recordatorios.")
    else:
        st.warning("Hay equipos próximos a calibrar dentro de los siguientes 30 días.")

    tabla_recordatorios = recordatorios.rename(columns={
        "id": "ID",
        "nombre": "Equipo",
        "tipo": "Tipo",
        "ubicacion": "Ubicación",
        "responsable": "Responsable",
        "proxima_calibracion": "Próxima calibración",
        "dias_restantes": "Días restantes",
        "estado_recordatorio": "Estado",
    })
    st.dataframe(tabla_recordatorios, use_container_width=True, hide_index=True)
    st.download_button(
        "Descargar recordatorios Word",
        data=generar_reporte_recordatorios_docx(tabla_recordatorios),
        file_name=f"recordatorios_calibracion_{date.today()}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
    )

st.divider()

st.markdown("<h3 style='color:#0b2c40; font-size:24px;'>Módulos disponibles</h3>", unsafe_allow_html=True)

col_a, col_b, col_c = st.columns(3)

with col_a:
    st.markdown("""
    <div style="background-color:#FFFFFF; border:1px solid #2dc197;
    border-radius:10px; padding:20px 24px; margin-bottom:8px; min-height:120px;">
        <p style="color:#1469aa; font-weight:700; font-size:17px; margin:0 0 8px 0;">📋 Ficha de equipos</p>
        <p style="color:#15795a; font-size:14px; margin:0;">
        Registro completo de instrumentos, historial de ubicación y especificaciones técnicas.</p>
    </div>
    """, unsafe_allow_html=True)
    if st.button("Gestionar equipos →", use_container_width=True, key="btn_ficha"):
        st.switch_page("pages/2_Ficha_de_equipos.py")

with col_b:
    st.markdown("""
    <div style="background-color:#FFFFFF; border:1px solid #2dc197;
    border-radius:10px; padding:20px 24px; margin-bottom:8px; min-height:120px;">
        <p style="color:#1469aa; font-weight:700; font-size:17px; margin:0 0 8px 0;">🔧 Control de mantenimientos</p>
        <p style="color:#15795a; font-size:14px; margin:0;">
        Programación y seguimiento de mantenimientos preventivos y correctivos.</p>
    </div>
    """, unsafe_allow_html=True)
    if st.button("Ver mantenimientos →", use_container_width=True, key="btn_mant"):
        st.switch_page("pages/3_Control_de_mantenimientos.py")

with col_c:
    st.markdown("""
    <div style="background-color:#FFFFFF; border:1px solid #2dc197;
    border-radius:10px; padding:20px 24px; margin-bottom:8px; min-height:120px;">
        <p style="color:#1469aa; font-weight:700; font-size:17px; margin:0 0 8px 0;">📈 Control de calibraciones</p>
        <p style="color:#15795a; font-size:14px; margin:0;">
        Gestión de incertidumbres, errores y resultados de calibración.</p>
    </div>
    """, unsafe_allow_html=True)
    if st.button("Gestionar calibraciones →", use_container_width=True, key="btn_cal"):
        st.switch_page("pages/4_Control_de_calibraciones.py")

col_d, col_e = st.columns(2)

with col_d:
    st.markdown("""
    <div style="background-color:#FFFFFF; border:1px solid #2dc197;
    border-radius:10px; padding:20px 24px; margin-bottom:8px; min-height:100px;">
        <p style="color:#1469aa; font-weight:700; font-size:17px; margin:0 0 8px 0;">✅ Condiciones de uso</p>
        <p style="color:#15795a; font-size:14px; margin:0;">
        Requisitos ambientales, operativos y de trazabilidad por instrumento.</p>
    </div>
    """, unsafe_allow_html=True)
    if st.button("Ver condiciones →", use_container_width=True, key="btn_cond"):
        st.switch_page("pages/5_Condiciones_de_uso.py")

with col_e:
    st.markdown("""
    <div style="background-color:#e8f5f0; border:1px solid #2dc197;
    border-radius:10px; padding:20px 24px; margin-bottom:8px; min-height:100px;">
        <p style="color:#063d7d; font-weight:700; font-size:17px; margin:0 0 8px 0;">ℹ️ Sobre la plataforma</p>
        <p style="color:#0a453c; font-size:14px; margin:0;">
        MetriCore cumple con los requisitos de la norma ISO/IEC 17025 para el 
        aseguramiento del control metrológico de instrumentos de medición.</p>
    </div>
    """, unsafe_allow_html=True)

st.divider()

st.markdown("""
<div style="background-color:#FFFFFF; border:1px solid #2dc197; border-left:5px solid #1469aa;
border-radius:10px; padding:16px 24px;">
    <p style="margin:0; color:#0b2c40; font-size:15px;">🔒 <strong>Plan Premium</strong> — Las funciones de 
    Regla de decisión e Intervalos de calibración requieren activar el Plan Premium. 
    Puedes hacerlo desde la sección <strong>Acceso a Plan Premium</strong> en el menú lateral.</p>
</div>
""", unsafe_allow_html=True)

st.divider()
st.caption("MetriCore v1.0 · SQLite · Streamlit · ISO/IEC 17025")
