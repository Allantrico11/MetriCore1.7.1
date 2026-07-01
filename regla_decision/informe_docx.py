"""
=============================================================================
Módulo generador de informes Word - Declaración de Conformidad
Basado en ISO/IEC 17025:2017 cláusula 7.8.6
=============================================================================

Genera un documento .docx profesional, listo para entregar al cliente, con:
  - Encabezado con datos del laboratorio/instrumento
  - Especificación aplicada
  - Regla de decisión utilizada
  - Tabla de resultados con declaración de conformidad y riesgos
  - Gráfico de zonas de conformidad incrustado
  - Declaración final según ISO/IEC 17025
"""

from __future__ import annotations
import io
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .logica_rd import INF


COLOR_VERDE = RGBColor(0x15, 0x92, 0x4A)
COLOR_AMARILLO = RGBColor(0xB8, 0x86, 0x0B)
COLOR_ROJO = RGBColor(0xC0, 0x39, 0x2B)
COLOR_AZUL = RGBColor(0x0B, 0x53, 0x94)
COLOR_GRIS_TEXTO = RGBColor(0x44, 0x44, 0x44)

_COLOR_POR_ICONO = {
    "✅": COLOR_VERDE,
    "⚠️": COLOR_AMARILLO,
    "❌": COLOR_ROJO,
}

COLOR_DURAZNO = RGBColor(0xCB, 0x4A, 0x1A)   # naranja oscuro para no-pasa-cond

_FILL_POR_ICONO = {
    "✅": "DCF5E3",
    "⚠️": "FCEFD0",
    "❌": "FBE0DC",
}

_FILL_POR_DECISION = {
    "Pasa":                     "DCF5E3",   # verde claro
    "Pasa condicionalmente":    "FCEFD0",   # naranja muy claro
    "No pasa condicionalmente": "FAD5BB",   # durazno/naranja más saturado
    "No pasa":                  "FBE0DC",   # rojo claro
}

_COLOR_POR_DECISION = {
    "Pasa":                     COLOR_VERDE,
    "Pasa condicionalmente":    COLOR_AMARILLO,
    "No pasa condicionalmente": COLOR_DURAZNO,
    "No pasa":                  COLOR_ROJO,
}


def _set_cell_background(cell, hex_color: str):
    """Aplica color de fondo a una celda de tabla."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def _set_cell_borders(cell):
    """Bordes finos grises para una celda."""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:color'), 'CCCCCC')
        borders.append(el)
    tc_pr.append(borders)


def _add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x0B, 0x53, 0x94)
    return h


def _add_kv_table(doc, pares, col1_width=Cm(5.5), col2_width=Cm(9.0)):
    """Tabla simple de dos columnas (campo / valor)."""
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for campo, valor in pares:
        row = table.add_row()
        row.cells[0].width = col1_width
        row.cells[1].width = col2_width
        p0 = row.cells[0].paragraphs[0]
        run0 = p0.add_run(str(campo))
        run0.bold = True
        run0.font.size = Pt(10.5)
        _set_cell_background(row.cells[0], "EFF4FA")
        _set_cell_borders(row.cells[0])

        p1 = row.cells[1].paragraphs[0]
        run1 = p1.add_run(str(valor))
        run1.font.size = Pt(10.5)
        _set_cell_borders(row.cells[1])
    return table


def _fmt_num(x, unidad="", decimales=5):
    if x is None:
        return "—"
    if x == INF or x == -INF:
        return "—"
    try:
        return f"{x:.{decimales}g} {unidad}".strip()
    except Exception:
        return str(x)


def _es_valido(x) -> bool:
    """True si x es un número válido (no None y no NaN)."""
    if x is None:
        return False
    try:
        return not (isinstance(x, float) and x != x)  # x != x detecta NaN sin depender de numpy
    except Exception:
        return True


def _fmt_pct_ref(valor: float, texto: str | None = None) -> str:
    if texto:
        return texto
    if valor <= 0.0001:
        return "1 ppm"
    return f"{valor:.4g}%"


def _riesgo_texto(fila: dict, cfg: dict | None = None) -> str:
    # Preferir el campo 'riesgo' precalculado (viene de app.py con el texto correcto)
    if fila.get("riesgo"):
        return fila["riesgo"]
    # Fallback: calcular desde valores numéricos
    decision = fila.get("decision", "")
    if decision == "Pasa":
        if cfg and cfg.get("regla") != "aceptacion_simple":
            return f"PFA ≤ {_fmt_pct_ref(cfg.get('riesgo_ref_pct', 2.5), cfg.get('riesgo_ref_text'))}"
        return "PFA ≤ 50%"
    elif decision == "No pasa":
        if cfg and cfg.get("regla") != "aceptacion_simple":
            return f"PFR ≤ {_fmt_pct_ref(cfg.get('riesgo_ref_pct', 2.5), cfg.get('riesgo_ref_text'))}"
        return "PFR ≤ 50%"
    pfa_min, pfa_max = fila.get("PFA_min"), fila.get("PFA_max")
    pfr_min, pfr_max = fila.get("PFR_min"), fila.get("PFR_max")
    partes = []
    if _es_valido(pfa_max):
        if _es_valido(pfa_min) and pfa_min > 0:
            partes.append(f"PFA entre {pfa_min:.2f}% y {pfa_max:.2f}%")
        else:
            partes.append(f"PFA ≤ {pfa_max:.2f}%")
    if _es_valido(pfr_max):
        if _es_valido(pfr_min) and pfr_min > 0:
            partes.append(f"PFR entre {pfr_min:.2f}% y {pfr_max:.2f}%")
        else:
            partes.append(f"PFR ≤ {pfr_max:.2f}%")
    return " · ".join(partes) if partes else "—"


def generar_informe_docx(cfg: dict, filas: list[dict], grafico_png_bytes: bytes | None,
                          grafico_individual_bytes: dict | None = None) -> bytes:
    """
    Genera el informe de conformidad en formato Word.

    cfg: diccionario con la configuración (nombre, unidad, nominal, LSL, USL,
         tipo_limite, regla, w, multiplicador, k, responsable, cliente, etc.)
    filas: lista de diccionarios, uno por cada punto de medición evaluado,
           con llaves: etiqueta, valor_medido, U, decision, icono, zona,
           PFA_min, PFA_max, PFR_min, PFR_max, AL_inf, AL_sup
    grafico_png_bytes: bytes PNG del gráfico comparativo (opcional)

    Devuelve los bytes del archivo .docx generado.
    """
    doc = Document()

    # Configuración de página y estilo base
    section = doc.sections[0]
    section.page_width = Cm(21.59)
    section.page_height = Cm(27.94)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)

    # ---------------------------------------------------------------------
    # PORTADA / ENCABEZADO
    # ---------------------------------------------------------------------
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run("INFORME DE DECLARACIÓN DE CONFORMIDAD")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = COLOR_AZUL

    subt = doc.add_paragraph()
    subt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subt.add_run("Regla de Decisión aplicada según ISO/IEC 17025:2017 (Cláusula 7.8.6)\nILAC G8:2009 · ISO 14253-1:2017")
    run2.italic = True
    run2.font.size = Pt(10)
    run2.font.color.rgb = COLOR_GRIS_TEXTO

    linea = doc.add_paragraph()
    p_border = OxmlElement('w:pPr')
    bottom = OxmlElement('w:pBdr')
    bdr = OxmlElement('w:bottom')
    bdr.set(qn('w:val'), 'single')
    bdr.set(qn('w:sz'), '12')
    bdr.set(qn('w:color'), '0B5394')
    bottom.append(bdr)
    p_border.append(bottom)
    linea._p.get_or_add_pPr().append(bottom)

    doc.add_paragraph()

    ahora = datetime.now().strftime("%d/%m/%Y %H:%M")
    _add_kv_table(doc, [
        ("Instrumento / Ítem evaluado", cfg.get('nombre', '—')),
        ("Cliente", cfg.get('cliente', '—') or "—"),
        ("Responsable del análisis", cfg.get('responsable', '—') or "—"),
        ("Fecha del informe", ahora),
        ("Unidad de medida", cfg.get('unidad', '—')),
    ])

    doc.add_paragraph()

    # ---------------------------------------------------------------------
    # 1. ESPECIFICACIÓN APLICADA
    # ---------------------------------------------------------------------
    _add_heading(doc, "1. Especificación Aplicada", level=2)
    LSL = cfg.get('LSL')
    USL = cfg.get('USL')
    unidad = cfg.get('unidad', '')

    filas_spec = [("Tipo de límite", cfg.get('tipo_limite', '—'))]
    if LSL is not None and LSL > -INF:
        filas_spec.append(("Límite inferior de tolerancia (LSL)", _fmt_num(LSL, unidad)))
    if USL is not None and USL < INF:
        filas_spec.append(("Límite superior de tolerancia (USL)", _fmt_num(USL, unidad)))
    if cfg.get('nominal') is not None:
        filas_spec.append(("Valor nominal", _fmt_num(cfg.get('nominal'), unidad)))
    if cfg.get('tolerancia') is not None:
        filas_spec.append(("Tolerancia (±)", _fmt_num(cfg.get('tolerancia'), unidad)))
    _add_kv_table(doc, filas_spec)

    doc.add_paragraph()

    # ---------------------------------------------------------------------
    # 2. REGLA DE DECISIÓN SELECCIONADA
    # ---------------------------------------------------------------------
    _add_heading(doc, "2. Regla de Decisión Seleccionada", level=2)

    nombre_regla = cfg.get('regla_nombre', cfg.get('regla', '—'))
    w_val = cfg.get('w')
    filas_regla = [("Regla de decisión", nombre_regla)]
    if w_val is not None:
        filas_regla.append(("Zona de seguridad (w)", _fmt_num(w_val, unidad)))
    if cfg.get('nivel_confianza_nombre'):
        filas_regla.append(("Nivel de confianza / criterio", cfg.get('nivel_confianza_nombre')))
    if cfg.get('k'):
        filas_regla.append(("Factor de cobertura (k)", str(cfg.get('k'))))
    _add_kv_table(doc, filas_regla)

    tipo_decl = ("No Binaria (Pasa / Pasa condicionalmente / "
                 "No pasa condicionalmente / No pasa)") if cfg.get('es_no_binaria') else \
                "Binaria (Pasa / No pasa)"
    p = doc.add_paragraph()
    p.add_run("Tipo de declaración: ").bold = True
    p.add_run(tipo_decl)

    doc.add_paragraph()

    # ---------------------------------------------------------------------
    # 3. RESULTADOS Y DECLARACIÓN DE CONFORMIDAD
    # ---------------------------------------------------------------------
    _add_heading(doc, "3. Resultados y Declaración de Conformidad", level=2)

    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    encabezados = ["Ítem", "Valor medido", "U", "Declaración", "Zona", "Riesgo asociado"]
    anchos = [Cm(2.0), Cm(2.6), Cm(1.8), Cm(3.4), Cm(4.0), Cm(3.7)]
    for i, texto in enumerate(encabezados):
        hdr[i].width = anchos[i]
        para = hdr[i].paragraphs[0]
        run = para.add_run(texto)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_background(hdr[i], "0B5394")
        _set_cell_borders(hdr[i])
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for fila in filas:
        row = table.add_row().cells
        icono    = fila.get("icono", "")
        decision = fila.get("decision", "")
        color_fill = _FILL_POR_DECISION.get(decision, _FILL_POR_ICONO.get(icono, "FFFFFF"))
        color_font = _COLOR_POR_DECISION.get(decision, _COLOR_POR_ICONO.get(icono, RGBColor(0, 0, 0)))
        valores = [
            str(fila.get("etiqueta", "")),
            _fmt_num(fila.get("valor_medido"), unidad),
            _fmt_num(fila.get("U"), unidad),
            fila.get("decision", ""),
            fila.get("zona", ""),
            _riesgo_texto(fila, cfg),
        ]
        for i, texto in enumerate(valores):
            row[i].width = anchos[i]
            para = row[i].paragraphs[0]
            run = para.add_run(texto)
            run.font.size = Pt(9.3)
            if i == 3:
                run.bold = True
                run.font.color.rgb = color_font
            _set_cell_background(row[i], color_fill)
            _set_cell_borders(row[i])

    doc.add_paragraph()

    # ---------------------------------------------------------------------
    # 4. GRÁFICO DE ZONAS DE CONFORMIDAD
    # ---------------------------------------------------------------------
    if grafico_png_bytes:
        _add_heading(doc, "4. Gráfico de Zonas de Conformidad", level=2)
        img_stream = io.BytesIO(grafico_png_bytes)
        doc.add_picture(img_stream, width=Cm(16.5))
        last_p = doc.paragraphs[-1]
        last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

    # ---------------------------------------------------------------------
    # 5. RESUMEN GLOBAL
    # ---------------------------------------------------------------------
    n_pasa          = sum(1 for f in filas if f.get("decision") == "Pasa")
    n_pasa_cond     = sum(1 for f in filas if f.get("decision") == "Pasa condicionalmente")
    n_no_pasa_cond  = sum(1 for f in filas if f.get("decision") == "No pasa condicionalmente")
    n_no_pasa       = sum(1 for f in filas if f.get("decision") == "No pasa")

    siguiente_heading = "5. Resumen Global" if grafico_png_bytes else "4. Resumen Global"
    _add_heading(doc, siguiente_heading, level=2)
    _add_kv_table(doc, [
        ("Total de mediciones evaluadas", str(len(filas))),
        ("Pasan",                         str(n_pasa)),
        ("Pasan condicionalmente",        str(n_pasa_cond)),
        ("No pasan condicionalmente",     str(n_no_pasa_cond)),
        ("No pasan",                      str(n_no_pasa)),
    ])

    doc.add_paragraph()

    # ---------------------------------------------------------------------
    # 6. DECLARACIÓN FINAL
    # ---------------------------------------------------------------------
    siguiente_heading2 = "6. Declaración Final" if grafico_png_bytes else "5. Declaración Final"
    _add_heading(doc, siguiente_heading2, level=2)
    p = doc.add_paragraph(
        f"Los resultados de medición presentados en este informe fueron comparados contra "
        f"los límites de especificación indicados, aplicando la regla de decisión "
        f"\u201c{nombre_regla}\u201d, conforme a los lineamientos de ISO/IEC 17025:2017 "
        f"(cláusula 7.8.6), ILAC G8:2009 e ISO 14253-1:2017."
    )
    p2 = doc.add_paragraph(
        "La regla de decisión aplicada fue acordada con el cliente o establecida por la "
        "organización, y se informa explícitamente en este documento conforme al requisito "
        "normativo correspondiente."
    )

    doc.add_paragraph()
    firma = doc.add_paragraph()
    firma.add_run("_" * 40 + "\n").font.size = Pt(10)
    firma.add_run(cfg.get('responsable', 'Responsable Técnico') or "Responsable Técnico").bold = True

    # ---------------------------------------------------------------------
    # Guardar en memoria
    # ---------------------------------------------------------------------
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
