"""
9_Intervalos_de_calibracion.py — Módulo de Intervalos de Calibración
======================================================================
Implementa los 3 métodos del ILAC-G24 / OIML D10:2007:

  MÉTODO 1 — Escalera (Error medio):
    Compara el error promedio del instrumento contra el 80% del EMP.
    Si está dentro → amplía IC 50%. Si está fuera → reduce IC 50%.
    Si supera el EMP → "Se debe ajustar y calibrar el equipo".
    NO aplica si el equipo fue ajustado recientemente.

  MÉTODO 2 — Escalera (Error con incertidumbre):
    Igual que el Método 1 pero compara |error| + U (incertidumbre
    expandida) en lugar del error solo. Más conservador.

  MÉTODO 3 — Cartas de control:
    Analiza el historial de varias calibraciones para calcular la
    tasa de deriva (cuánto cambia el error por año). Luego estima
    cuánto tiempo falta para que el error alcance el 100% del EMP.
    El IC final es el mínimo entre todos los puntos críticos.
    Usa 100% del EMP como límite de control.

Mejoras implementadas:
  - Validación de datos antes de calcular (evita divisiones por cero)
  - IC acotado: no supera el intervalo del fabricante ni baja de 3 meses
  - Contador de ajustes consecutivos del IC (patrón de comportamiento)
  - Gráfica por cada punto crítico en Cartas de control (selector)
  - Guardado del resultado en la base de datos SQLite
  - Intervalo mostrado en años, meses y días
  - Resumen del último cálculo IC al seleccionar un equipo
  - Confirmación antes de sobreescribir un resultado guardado

Integración con MetriCore:
  - Lee equipos y calibraciones desde SQLite (database.py)
  - Guarda resultados en tabla `resultados_ic` (nueva)
  - Requiere Plan Premium activo (st.session_state.premium_activo)

Paleta de colores:
  #23c057  #15924a  #063d7d  #238d93
  #0a453c  #1469aa  #15795a  #2dc197  #0b2c40
"""

import streamlit as st
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import io
from datetime import datetime, date, timedelta
from typing import Optional
import sys, os

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from database import get_equipos, get_calibraciones, get_connection, init_db
from utils import aplicar_estilos, mostrar_errores_validacion, validar_unidad

st.set_page_config(page_title="Intervalos de Calibración — MetriCore", layout="wide")

FECHA_MIN = date(1950, 1, 1)
FECHA_MAX = date(2150, 12, 31)
EMP_LABEL = "Error Máximo Permitido (EMP)"

# ── Paleta de colores ─────────────────────────────────────────────────────────
COLOR = {
    "verde":       "#23c057",
    "verde_osc":   "#15924a",
    "azul_marino": "#063d7d",
    "teal":        "#238d93",
    "verde_prof":  "#0a453c",
    "azul_med":    "#1469aa",
    "verde_menta": "#15795a",
    "aqua":        "#2dc197",
    "azul_noche":  "#0b2c40",
}


# ══════════════════════════════════════════════════════════════════════════════
# BASE DE DATOS — TABLA DE RESULTADOS IC
# ══════════════════════════════════════════════════════════════════════════════

def init_tabla_resultados_ic():
    """
    Crea la tabla `resultados_ic` si no existe.

    Esta tabla guarda cada cálculo de intervalo que el usuario confirma.
    Es independiente de la tabla `calibraciones` — esta tabla es propia
    del módulo de IC y registra los RESULTADOS del cálculo, no los datos
    de entrada del certificado de calibración.

    Columnas:
      equipo_id            : ID del equipo evaluado
      fecha_calculo        : cuándo se corrió el cálculo
      metodo               : "Escalera Error Medio", "Escalera Error U" o "Cartas de Control"
      emp                  : EMP usado en el cálculo
      unidad               : unidad de medida del instrumento
      intervalo_anios      : intervalo recomendado en años
      fecha_proxima_cal    : fecha exacta de próxima calibración
      recomendacion        : texto con la recomendación del sistema
      ajustes_consecutivos : cuántas veces seguidas se amplió o redujo el IC

    ⚠️  COORDINACIÓN: Esta tabla la crea y gestiona únicamente el módulo IC.
        El Dashboard puede leerla para mostrar alertas de calibración próxima
        usando: SELECT equipo_id, fecha_proxima_cal FROM resultados_ic
                WHERE fecha_calculo = (SELECT MAX(fecha_calculo) FROM resultados_ic r2
                WHERE r2.equipo_id = resultados_ic.equipo_id)
    """
    init_db()


def calcular_ajustes_consecutivos(historico: list[float], intervalo_actual: float) -> int:
    """
    Devuelve la cantidad de ajustes consecutivos en la misma dirección.

    Positivo = ampliaciones consecutivas, negativo = reducciones consecutivas,
    cero = sin patrón claro o intervalo igual al anterior.
    """
    if not historico or intervalo_actual is None:
        return 0

    anterior = historico[0]
    if intervalo_actual == anterior:
        return 0

    signo_actual = 1 if intervalo_actual > anterior else -1
    consecutivos = signo_actual

    for actual, previo in zip(historico, historico[1:]):
        if actual == previo:
            break
        signo = 1 if actual > previo else -1
        if signo != signo_actual:
            break
        consecutivos += signo_actual

    return consecutivos


def guardar_resultado_ic(equipo_id: str, metodo: str, emp: float, unidad: str,
                          intervalo_anios: float, fecha_proxima_cal: str,
                          recomendacion: str) -> bool:
    """
    Guarda el resultado del cálculo de IC en la base de datos.

    Además:
      1. Cuenta cuántos resultados consecutivos anteriores ampliaron o
         redujeron el IC en la misma dirección (para el contador de ajustes).
      2. Actualiza el campo `proxima_calibracion` en la tabla `equipos`,
         para que el Dashboard y otros módulos vean la fecha actualizada.

    Retorna True si se guardó correctamente, False si hubo error.
    """
    try:
        conn = get_connection()
        c    = conn.cursor()

        # Calcular ajustes consecutivos en la misma dirección
        c.execute("""
            SELECT intervalo_anios FROM resultados_ic
            WHERE equipo_id = ?
            ORDER BY id DESC
            LIMIT 5
        """, (equipo_id,))
        historico = [row[0] for row in c.fetchall() if row[0] is not None]

        ajustes_consecutivos = calcular_ajustes_consecutivos(historico, intervalo_anios)

        # Insertar resultado
        c.execute("""
            INSERT INTO resultados_ic
            (equipo_id, fecha_calculo, metodo, emp, unidad,
             intervalo_anios, fecha_proxima_cal, recomendacion, ajustes_consecutivos)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            equipo_id,
            date.today().strftime("%Y-%m-%d"),
            metodo, emp, unidad,
            intervalo_anios,
            fecha_proxima_cal,
            recomendacion,
            ajustes_consecutivos
        ))

        # Actualizar proxima_calibracion en tabla equipos para que lo vean otros módulos
        # ⚠️  COORDINACIÓN: Este campo lo lee el Dashboard para alertas.
        c.execute("""
            UPDATE equipos SET proxima_calibracion = ?
            WHERE id = ?
        """, (fecha_proxima_cal, equipo_id))

        conn.commit()
        conn.close()
        return True

    except Exception as e:
        st.error(f"Error al guardar en base de datos: {e}")
        return False


def get_historial_ic(equipo_id: str) -> pd.DataFrame:
    """
    Retorna el historial de cálculos de IC para un equipo específico,
    ordenado del más reciente al más antiguo.
    """
    try:
        conn = get_connection()
        df = pd.read_sql_query("""
            SELECT fecha_calculo, metodo, emp, unidad,
                   intervalo_anios, fecha_proxima_cal,
                   recomendacion, ajustes_consecutivos
            FROM resultados_ic
            WHERE equipo_id = ?
            ORDER BY id DESC
        """, conn, params=(equipo_id,))
        conn.close()
        return df
    except Exception:
        return pd.DataFrame()


# ══════════════════════════════════════════════════════════════════════════════
# FUNCIONES DE UTILIDAD
# ══════════════════════════════════════════════════════════════════════════════

def calcular_intervalo_anios(fecha_ant: date, fecha_act: date) -> float:
    """Calcula el intervalo en años entre dos fechas usando 365.25."""
    return (fecha_act - fecha_ant).days / 365.25


def fecha_desde_intervalo(fecha_base: date, intervalo_anios: float) -> date:
    """Calcula la fecha futura sumando un intervalo en años."""
    return fecha_base + timedelta(days=int(intervalo_anios * 365.25))


def formatear_intervalo(intervalo_anios: float) -> str:
    """
    Convierte un intervalo en años a texto con años, meses y días.

    Proceso:
      1. Convierte años → días totales (× 365.25)
      2. Extrae años completos (// 365)
      3. Del resto extrae meses (// 30)
      4. Lo que queda son días sueltos

    Ejemplos:
      1.50 → "1 año, 6 meses y 2 días"
      0.50 → "6 meses y 1 día"
      0.25 → "3 meses y 2 días"
    """
    dias_totales = int(intervalo_anios * 365.25)
    anios  = dias_totales // 365
    resto  = dias_totales % 365
    meses  = resto // 30
    dias   = resto % 30

    partes = []
    if anios > 0:
        partes.append(f"{anios} {'año' if anios == 1 else 'años'}")
    if meses > 0:
        partes.append(f"{meses} {'mes' if meses == 1 else 'meses'}")
    if dias > 0:
        partes.append(f"{dias} {'día' if dias == 1 else 'días'}")

    if not partes:
        return "0 días"
    if len(partes) == 1:
        return partes[0]
    return ", ".join(partes[:-1]) + f" y {partes[-1]}"


def acotar_intervalo(intervalo: float, intervalo_fabricante: float) -> tuple:
    """
    Aplica el límite mínimo al intervalo calculado y advierte cuando supera
    el intervalo recomendado por el fabricante.

    Según ILAC-G24, el IC no debería:
      - Advertir si supera el intervalo recomendado por el fabricante
      - Ser menor a 3 meses = 0.25 años (piso mínimo razonable)

    Retorna (intervalo_recomendado, mensaje_advertencia_o_None)
    """
    MINIMO_ANIOS = 0.25   # 3 meses
    mensaje = None

    if intervalo < MINIMO_ANIOS:
        mensaje = (f"El intervalo calculado ({intervalo:.2f} años) es menor al mínimo "
                   f"recomendado de 3 meses. Se ajusta a 3 meses.")
        intervalo = MINIMO_ANIOS

    if intervalo > intervalo_fabricante:
        mensaje = (f"El intervalo calculado ({intervalo:.2f} años) supera el recomendado "
                   f"por el fabricante ({intervalo_fabricante:.2f} años). "
                   f"Se mantiene como recomendado el intervalo calculado por el método, "
                   f"pero conviene revisar esta diferencia antes de aprobarlo.")

    return intervalo, mensaje


def validar_datos_escalera(error_medio: float, emp: float,
                            incertidumbre: float, fecha_ant: date,
                            fecha_act: date) -> list:
    """
    Valida los datos de entrada antes de calcular con los métodos de Escalera.
    Retorna lista de mensajes de error. Si está vacía, los datos son válidos.

    Validaciones:
      - EMP > 0 (no se puede dividir entre cero)
      - Fecha actual posterior a la anterior
      - Incertidumbre no negativa
      - El error ingresado tiene signo (puede ser negativo, es válido)
    """
    errores = []
    if emp <= 0:
        errores.append(f"El {EMP_LABEL} debe ser mayor que cero.")
    if fecha_act <= fecha_ant:
        errores.append("La fecha de calibración actual debe ser posterior a la anterior.")
    if incertidumbre < 0:
        errores.append("La incertidumbre expandida no puede ser negativa.")
    return errores


def validar_datos_cartas(fechas: list, emp: float) -> list:
    """
    Valida los datos de entrada antes de calcular con Cartas de control.

    Validaciones:
      - Al menos 2 fechas
      - Fechas en orden cronológico
      - EMP > 0
      - No hay fechas duplicadas (causaría división por cero en la regresión)
    """
    errores = []
    if emp <= 0:
        errores.append(f"El {EMP_LABEL} debe ser mayor que cero.")
    if len(fechas) < 2:
        errores.append("Se necesitan al menos 2 fechas de calibración.")
    if fechas != sorted(fechas):
        errores.append("Las fechas deben estar en orden cronológico.")
    if len(fechas) != len(set(fechas)):
        errores.append("No puede haber fechas duplicadas.")
    return errores


# ══════════════════════════════════════════════════════════════════════════════
# FUNCIONES DE CÁLCULO
# ══════════════════════════════════════════════════════════════════════════════

def escalera_error_medio(error_medio: float, emp: float,
                          intervalo_anterior: float, fue_ajustado: bool) -> dict:
    """
    MÉTODO 1 — Escalera usando el error medio.
    Referencia: ILAC-G24 sección 4.4.1

    Lógica:
      Límite de control = 80% del EMP
      - Si ajustado → no aplica
      - Si |error| ≤ límite → IC × 1.50 (ampliar 50%)
      - Si límite < |error| < EMP → IC × 0.50 (reducir 50%)
      - Si |error| ≥ EMP → ajuste y calibración inmediata
    """
    limite = 0.80 * emp

    if abs(error_medio) >= emp:
        return {
            "recomendacion": f"El error medio ({error_medio:.5f}) alcanza o supera el {EMP_LABEL} ({emp:.5f}). "
                             "Se debe ajustar y calibrar el equipo.",
            "intervalo_nuevo": None, "tipo": "bad"
        }
    if fue_ajustado:
        return {
            "recomendacion": "El equipo fue ajustado recientemente. "
                             "Use el intervalo recomendado por el fabricante.",
            "intervalo_nuevo": None, "tipo": "warn"
        }
    elif abs(error_medio) <= limite:
        return {
            "recomendacion": (f"El error medio ({error_medio:.5f}) está dentro del "
                              f"límite de control ({limite:.5f}). "
                              f"Se amplía el intervalo en un 50%."),
            "intervalo_nuevo": intervalo_anterior * 1.50, "tipo": "ok"
        }
    else:
        return {
            "recomendacion": (f"El error medio ({error_medio:.5f}) supera el "
                              f"límite de control ({limite:.5f}). "
                              f"Se reduce el intervalo en un 50%."),
            "intervalo_nuevo": intervalo_anterior * 0.50, "tipo": "warn"
        }


def escalera_error_incertidumbre(error: float, incertidumbre: float, emp: float,
                                  intervalo_anterior: float, fue_ajustado: bool) -> dict:
    """
    MÉTODO 2 — Escalera con error ± incertidumbre expandida.
    Referencia: ILAC-G24 sección 4.4.2

    Diferencia clave respecto al Método 1:
      valor_comparar = |error| + U  (peor caso posible)
    Esto es más conservador porque reconoce la duda asociada al error medido.
    """
    limite      = 0.80 * emp
    error_con_u = abs(error) + incertidumbre

    if error_con_u >= emp:
        return {
            "recomendacion": (f"El error ± U ({error_con_u:.5f}) alcanza o supera el {EMP_LABEL} ({emp:.5f}). "
                              "Se debe ajustar y calibrar el equipo."),
            "intervalo_nuevo": None, "tipo": "bad"
        }
    if fue_ajustado:
        return {
            "recomendacion": "El equipo fue ajustado recientemente. "
                             "Use el intervalo recomendado por el fabricante.",
            "intervalo_nuevo": None, "tipo": "warn"
        }
    elif error_con_u <= limite:
        return {
            "recomendacion": (f"El error ± U ({error_con_u:.5f}) está dentro del "
                              f"límite de control ({limite:.5f}). "
                              f"Se amplía el intervalo en un 50%."),
            "intervalo_nuevo": intervalo_anterior * 1.50, "tipo": "ok"
        }
    else:
        return {
            "recomendacion": (f"El error ± U ({error_con_u:.5f}) supera el "
                              f"límite de control ({limite:.5f}). "
                              f"Se reduce el intervalo en un 50%."),
            "intervalo_nuevo": intervalo_anterior * 0.50, "tipo": "warn"
        }


def cartas_control(fechas: list, errores: list, emp: float,
                   fue_ajustado: bool, fecha_ajuste: Optional[date] = None) -> dict:
    """
    MÉTODO 3 — Cartas de control (análisis de deriva).
    Referencia: ILAC-G24 sección 4.4.3

    Lógica paso a paso:
      1. Convierte fechas a tiempos en años desde la primera fecha
      2. Si hubo ajuste: filtra datos post-ajuste y agrega punto (t=0, error=0)
      3. Regresión lineal: error(t) = pendiente × t + intercepto
         → pendiente = deriva anual
      4. Error actual = valor de la recta en el último tiempo
      5. Margen = 100% EMP − error actual
         ⚠️ Usa 100% del EMP, NO 80% como en Escalera
      6. IC = margen / deriva_anual
    """
    fecha_ref     = fechas[0]
    tiempos_anios = [(f - fecha_ref).days / 365.25 for f in fechas]

    if fue_ajustado and fecha_ajuste:
        datos_post = [(t, e) for t, e, f in zip(tiempos_anios, errores, fechas)
                      if f >= fecha_ajuste]
        if len(datos_post) < 2:
            return {"recomendacion": "Con ajuste reciente se necesitan al menos 2 calibraciones post-ajuste.",
                    "intervalo_nuevo": None, "tipo": "warn",
                    "deriva_anual": None, "datos_grafica": None}
        tiempos_anios = [0.0] + [d[0] for d in datos_post]
        errores       = [0.0] + [d[1] for d in datos_post]

    t = np.array(tiempos_anios)
    e = np.array(errores)

    pendiente, intercepto = np.polyfit(t, e, 1)
    deriva_anual          = abs(pendiente)
    limite_control        = 1.00 * emp   # ⚠️ 100% del EMP

    datos_grafica = {
        "tiempos": list(t), "errores": list(e),
        "pendiente": pendiente, "intercepto": intercepto,
        "limite_control": limite_control, "emp": emp
    }

    if deriva_anual == 0:
        return {"recomendacion": "No se detectó deriva. Use el intervalo del fabricante.",
                "intervalo_nuevo": None, "tipo": "warn",
                "deriva_anual": 0, "datos_grafica": datos_grafica}

    error_actual = abs(intercepto + pendiente * t[-1])
    margen       = limite_control - error_actual

    if margen <= 0:
        return {"recomendacion": (f"La deriva ya supera el 100% del {EMP_LABEL}. "
                                  "Se debe ajustar y calibrar el equipo."),
                "intervalo_nuevo": None, "tipo": "bad",
                "deriva_anual": deriva_anual, "datos_grafica": datos_grafica}

    intervalo = margen / deriva_anual
    return {
        "recomendacion": (f"Deriva estimada: {deriva_anual:.5f}/año. "
                          f"El instrumento alcanzará el 100% del {EMP_LABEL} en {intervalo:.2f} años."),
        "intervalo_nuevo": intervalo,
        "tipo": "ok",
        "deriva_anual": deriva_anual,
        "datos_grafica": datos_grafica
    }


# ══════════════════════════════════════════════════════════════════════════════
# FUNCIONES DE GRÁFICAS
# ══════════════════════════════════════════════════════════════════════════════

def grafica_escalera(error_val: float, emp: float, unidad: str,
                     con_incertidumbre: bool = False, incertidumbre: float = 0.0):
    """
    Gráfica de evaluación para métodos de Escalera.
    Zona verde = dentro del 80% EMP. Zona amarilla = entre 80% y menor que EMP.
    Alcanzar o superar el EMP se muestra como condición crítica.
    """
    fig, ax = plt.subplots(figsize=(7, 5.2))
    fig.patch.set_facecolor("#0b2c40")
    ax.set_facecolor("#0b2c40")
    ax.set_position([0.12, 0.34, 0.82, 0.52])
    limite = 0.80 * emp

    ax.axhspan(-emp,    -limite, alpha=0.15, color="#f44336")
    ax.axhspan(-limite,  limite, alpha=0.15, color="#23c057")
    ax.axhspan( limite,  emp,   alpha=0.15, color="#f44336")

    ax.axhline(y= emp,    color="#f44336", ls="--", lw=1.2, label=f"+{EMP_LABEL} ({emp:.4f})")
    ax.axhline(y=-emp,    color="#f44336", ls="--", lw=1.2, label=f"−{EMP_LABEL} (−{emp:.4f})")
    ax.axhline(y= limite, color="#ff9800", ls=":",  lw=1.2, label=f"+80% {EMP_LABEL} ({limite:.4f})")
    ax.axhline(y=-limite, color="#ff9800", ls=":",  lw=1.2, label=f"−80% {EMP_LABEL} (−{limite:.4f})")
    ax.axhline(y=0, color="white", ls="-", lw=0.5, alpha=0.3)

    valor_evaluado = abs(error_val) + incertidumbre if con_incertidumbre else abs(error_val)
    color_punto = (COLOR["verde"] if valor_evaluado <= limite
                   else ("#ff9800" if valor_evaluado < emp else "#f44336"))

    if con_incertidumbre and incertidumbre > 0:
        ax.errorbar(0, error_val, yerr=incertidumbre, fmt="o", color=color_punto,
                    markersize=10, capsize=6, capthick=2, elinewidth=2, label="Error ± U")
    else:
        ax.scatter(0, error_val, color=color_punto, s=120, zorder=5, label="Error medio")

    ax.scatter(1.35, 0, alpha=0, s=1, label="_nolegend_")
    ax.set_xlim(-1, 1.65)
    ax.set_xticks([])
    ax.set_ylabel(f"Error ({unidad})", color="white")
    ax.tick_params(colors="white")
    ax.spines[:].set_color("#2dc19733")
    ax.set_title("Evaluación — Método de Escalera", color="white", fontsize=11, pad=10)
    leyenda = ax.legend(loc="center right", frameon=True, fontsize=8)
    leyenda.get_frame().set_facecolor("#ffffff")
    leyenda.get_frame().set_edgecolor("#2dc197")
    leyenda.get_frame().set_alpha(0.92)
    return fig


def grafica_cartas_control(datos: dict, unidad: str, titulo_punto: str = "Punto 1"):
    """
    Gráfica de deriva para Cartas de control.
    Límite al 100% del EMP. Línea azul = tendencia (regresión lineal).
    """
    if not datos:
        return None

    fig, ax = plt.subplots(figsize=(8, 5.4))
    fig.patch.set_facecolor("#0b2c40")
    ax.set_facecolor("#0b2c40")
    ax.set_position([0.10, 0.36, 0.86, 0.52])

    t   = np.array(datos["tiempos"])
    e   = np.array(datos["errores"])
    m   = datos["pendiente"]
    b   = datos["intercepto"]
    lc  = datos["limite_control"]
    emp = datos["emp"]

    t_max = max(t) if len(t) else 1.0
    if t_max <= 0:
        t_max = 1.0
    t_ext = np.linspace(0, t_max * 1.5, 200)

    ax.axhspan(-emp * 1.2, -lc, alpha=0.12, color="#f44336")
    ax.axhspan(-lc,         lc, alpha=0.12, color="#23c057")
    ax.axhspan( lc, emp * 1.2, alpha=0.12, color="#f44336")

    ax.axhline(y= emp, color="#f44336", ls="--", lw=1.2, label=f"+100% {EMP_LABEL} ({emp:.4f})")
    ax.axhline(y=-emp, color="#f44336", ls="--", lw=1.2, label=f"−100% {EMP_LABEL}")
    ax.plot(t_ext, m * t_ext + b, color=COLOR["aqua"], lw=1.8,
            ls="-.", label="Tendencia (deriva)")
    ax.scatter(t, e, color=COLOR["verde"], s=80, zorder=5, label="Calibraciones")
    ax.plot(t, e, color=COLOR["verde"], lw=1, alpha=0.5)
    ax.scatter(t_max * 1.65, 0, alpha=0, s=1, label="_nolegend_")
    ax.set_xlim(min(0, min(t)), t_max * 1.85)

    ax.set_xlabel("Tiempo desde primera calibración (años)", color="white")
    ax.set_ylabel(f"Error ({unidad})", color="white")
    ax.tick_params(colors="white")
    ax.spines[:].set_color("#2dc19733")
    ax.set_title(f"Cartas de Control — {titulo_punto} (límite: 100% {EMP_LABEL})",
                 color="white", fontsize=11, pad=10)
    leyenda = ax.legend(loc="center right", frameon=True, fontsize=8)
    leyenda.get_frame().set_facecolor("#ffffff")
    leyenda.get_frame().set_edgecolor("#2dc197")
    leyenda.get_frame().set_alpha(0.92)
    return fig


def leyenda_html(items: list[tuple[str, str]]) -> str:
    """Construye una leyenda externa sin HTML para evitar que Streamlit muestre código."""
    simbolos = {
        "#f44336": "Linea roja",
        "#ff9800": "Linea naranja",
        "#23c057": "Punto verde",
        COLOR["aqua"]: "Linea azul",
        COLOR["verde"]: "Punto verde",
    }
    lineas = ["**Leyenda del gráfico:**"]
    for color, texto in items:
        lineas.append(f"- {simbolos.get(color, 'Referencia')}: {texto}")
    return "\n".join(lineas)


def leyenda_escalera_html(emp: float, unidad: str, limite: float,
                          etiqueta_punto: str) -> str:
    return leyenda_html([
        ("#f44336", f"+/- {EMP_LABEL}: {emp:.4f} {unidad}".strip()),
        ("#ff9800", f"+/- 80% {EMP_LABEL}: {limite:.4f} {unidad}".strip()),
        ("#23c057", etiqueta_punto),
    ])


def leyenda_cartas_html(emp: float, unidad: str) -> str:
    return leyenda_html([
        ("#f44336", f"+/- 100% {EMP_LABEL}: {emp:.4f} {unidad}".strip()),
        (COLOR["aqua"], "Tendencia de deriva"),
        (COLOR["verde"], "Calibraciones registradas"),
    ])


def _doc_set_cell_background(cell, hex_color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _doc_set_cell_borders(cell, color: str = "CCCCCC"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), color)
        borders.append(el)
    tc_pr.append(borders)


def _doc_heading(doc: Document, text: str, level: int = 2):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x14, 0x69, 0xAA)
    return h


def _doc_kv_table(doc: Document, rows: list[tuple[str, str]], col1_width=Cm(5.7), col2_width=Cm(9.0)):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].width = col1_width
        cells[1].width = col2_width
        p0 = cells[0].paragraphs[0]
        r0 = p0.add_run(str(key))
        r0.bold = True
        r0.font.size = Pt(10.3)
        r0.font.color.rgb = RGBColor(0x0B, 0x2C, 0x40)
        _doc_set_cell_background(cells[0], "E0F0ED")
        _doc_set_cell_borders(cells[0])
        p1 = cells[1].paragraphs[0]
        r1 = p1.add_run(str(value))
        r1.font.size = Pt(10.3)
        r1.font.color.rgb = RGBColor(0x0B, 0x2C, 0x40)
        _doc_set_cell_borders(cells[1])
    return table


def _doc_style_header(cell, text: str):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(9.2)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    _doc_set_cell_background(cell, "1469AA")
    _doc_set_cell_borders(cell, "1469AA")


def generar_reporte_ic_docx(equipo: Optional[dict], unidad: str, emp: float,
                            met_usado: str, intervalo_r: float,
                            texto_ic: str, fecha_proxima: date,
                            recomendacion: str, msg_acotado: Optional[str],
                            resultados_puntos: Optional[list] = None) -> bytes:
    """Genera un reporte Word del cálculo de intervalo de calibración."""
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)

    banner = doc.add_table(rows=1, cols=1)
    banner.alignment = WD_TABLE_ALIGNMENT.CENTER
    banner_cell = banner.rows[0].cells[0]
    _doc_set_cell_background(banner_cell, "0B2C40")
    _doc_set_cell_borders(banner_cell, "0B2C40")
    p_banner = banner_cell.paragraphs[0]
    p_banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_banner = p_banner.add_run("MetriCore")
    r_banner.bold = True
    r_banner.font.size = Pt(12)
    r_banner.font.color.rgb = RGBColor(0xE0, 0xF0, 0xED)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("REPORTE DE INTERVALO DE CALIBRACIÓN")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x06, 0x3D, 0x7D)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("ILAC-G24 / OIML D10 · Gestión de intervalos de calibración")
    sub_run.italic = True
    sub_run.font.color.rgb = RGBColor(0x15, 0x79, 0x5A)

    doc.add_paragraph()
    _doc_heading(doc, "1. Equipo evaluado")
    _doc_kv_table(doc, [
        ("Equipo", equipo["nombre"] if equipo else "Ingreso manual"),
        ("ID", equipo["id"] if equipo else "—"),
        ("Unidad", unidad or "—"),
        ("Fecha de cálculo", date.today().strftime("%d/%m/%Y")),
    ])

    doc.add_paragraph()
    _doc_heading(doc, "2. Parámetros y método")
    _doc_kv_table(doc, [
        ("Método", met_usado),
        (f"{EMP_LABEL} / tolerancia usada", f"{emp:g} {unidad}".strip()),
        ("Intervalo recomendado", f"{intervalo_r:.2f} años ({texto_ic})"),
        ("Próxima calibración", fecha_proxima.strftime("%d/%m/%Y")),
    ])

    doc.add_paragraph()
    _doc_heading(doc, "3. Recomendación")
    rec_table = doc.add_table(rows=1, cols=1)
    rec_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    rec_cell = rec_table.rows[0].cells[0]
    _doc_set_cell_background(rec_cell, "DCF5E3")
    _doc_set_cell_borders(rec_cell, "2DC197")
    rec_p = rec_cell.paragraphs[0]
    rec_run = rec_p.add_run(recomendacion)
    rec_run.bold = True
    rec_run.font.color.rgb = RGBColor(0x0A, 0x45, 0x3C)

    if msg_acotado:
        doc.add_paragraph()
        warn_table = doc.add_table(rows=1, cols=1)
        warn_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        warn_cell = warn_table.rows[0].cells[0]
        _doc_set_cell_background(warn_cell, "FCEFD0")
        _doc_set_cell_borders(warn_cell, "B8860B")
        warn_p = warn_cell.paragraphs[0]
        warn_title = warn_p.add_run("Advertencia: ")
        warn_title.bold = True
        warn_title.font.color.rgb = RGBColor(0xB8, 0x86, 0x0B)
        warn_text = warn_p.add_run(msg_acotado)
        warn_text.font.color.rgb = RGBColor(0x0B, 0x2C, 0x40)

    if resultados_puntos:
        doc.add_paragraph()
        _doc_heading(doc, "4. Detalle por punto crítico")
        table = doc.add_table(rows=1, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        headers = ["Nominal", "Deriva/año", "IC (años)", "IC detalle", "Estado"]
        for idx, header in enumerate(headers):
            _doc_style_header(table.rows[0].cells[idx], header)
        for punto in resultados_puntos:
            ic_p = punto.get("intervalo_nuevo")
            row = table.add_row().cells
            row[0].text = str(punto.get("nominal", "—"))
            row[1].text = f"{punto.get('deriva_anual'):.5f}" if punto.get("deriva_anual") else "—"
            row[2].text = f"{ic_p:.2f}" if ic_p else "—"
            row[3].text = formatear_intervalo(ic_p) if ic_p else "—"
            row[4].text = punto.get("tipo", "—")
            for cell in row:
                _doc_set_cell_borders(cell)

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Generado por MetriCore · Reporte técnico")
    footer_run.italic = True
    footer_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

# ══════════════════════════════════════════════════════════════════════════════
# SECCIÓN DE RESULTADOS
# ══════════════════════════════════════════════════════════════════════════════

def mostrar_resultado(res: dict, fecha_base: date, met_usado: str,
                      equipo: Optional[dict], unidad: str, emp: float,
                      intervalo_fabricante: float,
                      resultados_puntos: Optional[list] = None):
    """
    Muestra la sección de resultados completa y uniforme para los 3 métodos.

    Incluye:
      - Caja de recomendación con color según severidad
      - Métricas: intervalo en años y desglosado en meses/días
      - Advertencia si el IC fue acotado (mínimo o máximo)
      - Contador de ajustes consecutivos con alerta si es ≥ 3
      - Selector de punto para ver su gráfica (Cartas de control)
      - Botón para guardar en base de datos
      - Exportar reporte en .txt
      - Historial de cálculos anteriores del equipo
    """
    # ── Acotar el intervalo ───────────────────────────────────────────────────
    intervalo_calculado = res["intervalo_nuevo"]
    intervalo_r         = intervalo_calculado
    msg_acotado         = None
    criterio_guardado   = "método"
    recomendacion_guardada = res["recomendacion"]

    if intervalo_r is not None:
        intervalo_r, msg_acotado = acotar_intervalo(intervalo_r, intervalo_fabricante)

    supera_fabricante = (
        intervalo_calculado is not None
        and intervalo_fabricante is not None
        and intervalo_calculado > intervalo_fabricante
    )

    if supera_fabricante:
        st.warning(
            "El intervalo calculado por el método supera el intervalo recomendado por "
            "el fabricante. Antes de guardar, seleccione qué criterio quedará en el registro."
        )
        criterio_opcion = st.radio(
            "Criterio para guardar en el registro:",
            [
                "Guardar intervalo calculado por el método",
                "Guardar intervalo recomendado por el fabricante",
            ],
            horizontal=False,
            key=f"ic_criterio_guardado_{equipo.get('id') if equipo else 'manual'}_{met_usado}",
        )
        if criterio_opcion == "Guardar intervalo recomendado por el fabricante":
            intervalo_r = intervalo_fabricante
            criterio_guardado = "fabricante"
            msg_acotado = (
                f"El intervalo calculado ({intervalo_calculado:.2f} años) supera el recomendado "
                f"por el fabricante ({intervalo_fabricante:.2f} años). "
                "Se seleccionó guardar el intervalo del fabricante como criterio del registro."
            )
        else:
            criterio_guardado = "método"
            msg_acotado = (
                f"El intervalo calculado ({intervalo_calculado:.2f} años) supera el recomendado "
                f"por el fabricante ({intervalo_fabricante:.2f} años). "
                "Se seleccionó guardar el intervalo calculado por el método como criterio del registro."
            )

        recomendacion_guardada = (
            f"{res['recomendacion']} Criterio guardado: intervalo del {criterio_guardado}."
        )

    # ── Caja de recomendación ─────────────────────────────────────────────────
    color_box  = {"ok": "#23c05722", "warn": "#ff980022", "bad": "#f4433622"}.get(res["tipo"], "#ff980022")
    border_box = {"ok": "#23c057",   "warn": "#ff9800",   "bad": "#f44336"  }.get(res["tipo"], "#ff9800")
    icono      = {"ok": "✅",        "warn": "⚠️",        "bad": "🚨"       }.get(res["tipo"], "ℹ️")

    st.markdown(f"""
    <div style="background:{color_box}; border-left:4px solid {border_box};
    border-radius:10px; padding:1rem 1.4rem; margin-top:0.8rem;">
        <strong>{icono} {res['recomendacion']}</strong>
    </div>
    """, unsafe_allow_html=True)

    if res.get("tipo") == "bad":
        st.error(f"Acción requerida: {res['recomendacion']}")

    if msg_acotado:
        st.warning(f"Advertencia sobre intervalo: {msg_acotado}")

    # ── Métricas ──────────────────────────────────────────────────────────────
    if intervalo_r is not None and fecha_base:
        fecha_proxima = fecha_desde_intervalo(fecha_base, intervalo_r)
        texto_detalle = formatear_intervalo(intervalo_r)

        col1, col2, col3 = st.columns(3)
        col1.metric("Intervalo recomendado", f"{intervalo_r:.2f} años")
        col2.metric("Equivale a",            texto_detalle)
        col3.metric("Próxima calibración",   fecha_proxima.strftime("%d/%m/%Y"))

    # ── Contador de ajustes consecutivos ─────────────────────────────────────
    if equipo and equipo.get("id"):
        hist_ic = get_historial_ic(equipo["id"])
        if not hist_ic.empty and len(hist_ic) >= 2:
            ultimo    = hist_ic.iloc[0]["intervalo_anios"]
            penultimo = hist_ic.iloc[1]["intervalo_anios"]
            if ultimo and penultimo and intervalo_r:
                # Detectar si el patrón actual continúa en la misma dirección
                subiendo = intervalo_r > ultimo > penultimo
                bajando  = intervalo_r < ultimo < penultimo
                n_consec = int(hist_ic.iloc[0]["ajustes_consecutivos"]) + 1 if (subiendo or bajando) else 0

                if n_consec >= 3:
                    direccion = "ampliar" if subiendo else "reducir"
                    st.info(
                        f"📊 El IC se ha {direccion}do **{n_consec} veces consecutivas**. "
                        f"Esto puede indicar un patrón de comportamiento del instrumento "
                        f"que vale la pena revisar con el proveedor de calibración."
                    )

    # ── Detalle por punto (Cartas de control) ─────────────────────────────────
    if resultados_puntos and len(resultados_puntos) > 0:
        st.markdown("**Detalle por punto crítico:**")
        filas = []
        for p in resultados_puntos:
            ic_p = p.get("intervalo_nuevo")
            if ic_p:
                ic_p, _ = acotar_intervalo(ic_p, intervalo_fabricante)
            filas.append({
                "Nominal":    p["nominal"],
                "Deriva/año": f"{p['deriva_anual']:.5f}" if p.get("deriva_anual") else "—",
                "IC (años)":  f"{ic_p:.2f}" if ic_p else "—",
                "IC detalle": formatear_intervalo(ic_p) if ic_p else "—",
                "Estado":     {"ok": "✅", "warn": "⚠️", "bad": "🚨"}.get(p["tipo"], "—")
            })
        st.dataframe(pd.DataFrame(filas), use_container_width=True, hide_index=True)

        # Selector de gráfica por punto
        st.markdown("**Gráfica por punto:**")
        nombres_puntos = [f"Punto {i+1} (nominal: {p['nominal']})"
                          for i, p in enumerate(resultados_puntos)]
        punto_sel = st.selectbox("Ver gráfica del punto:", nombres_puntos)
        idx_sel   = nombres_puntos.index(punto_sel)

        datos_g = resultados_puntos[idx_sel].get("datos_grafica")
        if datos_g:
            fig = grafica_cartas_control(datos_g, unidad, titulo_punto=punto_sel)
            if fig:
                st.pyplot(fig)
                plt.close(fig)
        else:
            st.caption("No hay datos de gráfica para este punto.")

    else:
        # Gráfica única (métodos de Escalera)
        fig_res = st.session_state.get("ic_fig")
        if fig_res:
            st.markdown("**Visualización:**")
            st.pyplot(fig_res)
            etiqueta_punto = (
                "Error ± U"
                if "incertidumbre" in str(met_usado).lower()
                else "Error medio"
            )
            plt.close(fig_res)

    # ── Guardar en base de datos ──────────────────────────────────────────────
    st.divider()
    st.subheader("5 · Guardar resultado")

    if equipo and equipo.get("id") and intervalo_r is not None and fecha_base:
        fecha_proxima     = fecha_desde_intervalo(fecha_base, intervalo_r)
        fecha_proxima_str = fecha_proxima.strftime("%Y-%m-%d")

        st.caption(
            "Guardar registra este cálculo en el historial del equipo y actualiza "
            "la fecha de próxima calibración para que el Dashboard pueda generar alertas."
        )

        # Verificar si ya existe un resultado guardado para mostrar confirmación
        hist_existente = get_historial_ic(equipo["id"])
        ya_tiene_resultado = not hist_existente.empty
        guardado_ok = st.session_state.get("ic_guardado_ok")
        guardado_actual = (
            isinstance(guardado_ok, dict)
            and guardado_ok.get("equipo_id") == equipo["id"]
        )

        if guardado_actual:
            st.success(guardado_ok.get("mensaje", "✅ Resultado guardado correctamente."))
            confirmar = False
        elif ya_tiene_resultado:
            # Mostrar el resultado anterior para que el usuario lo compare
            ult = hist_existente.iloc[0]
            st.warning(
                f"⚠️ Este equipo ya tiene un resultado IC guardado: "
                f"**{ult['intervalo_anios']:.2f} años** calculado el "
                f"**{ult['fecha_calculo']}** con el método **{ult['metodo']}**. "
                f"¿Desea guardar este nuevo cálculo como una nueva versión del historial?"
            )
            col_si, col_no = st.columns(2)
            confirmar = col_si.button("✅ Sí, guardar en historial", type="primary")
            cancelar  = col_no.button("❌ No, cancelar")

            if cancelar:
                st.info("Guardado cancelado. El historial anterior se conserva.")
                confirmar = False
        else:
            confirmar = st.button("💾 Guardar resultado en base de datos", type="primary")

        if confirmar:
            ok = guardar_resultado_ic(
                equipo_id        = equipo["id"],
                metodo           = met_usado,
                emp              = emp,
                unidad           = unidad,
                intervalo_anios  = intervalo_r,
                fecha_proxima_cal= fecha_proxima_str,
                recomendacion    = recomendacion_guardada
            )
            if ok:
                st.session_state["ic_guardado_ok"] = {
                    "equipo_id": equipo["id"],
                    "mensaje": (
                        f"✅ Resultado guardado para **{equipo['nombre']}**. "
                        f"Próxima calibración: {fecha_proxima.strftime('%d/%m/%Y')}."
                    ),
                }
                st.rerun()
            else:
                st.error("No se pudo guardar. Verifique que la base de datos existe.")

    elif not equipo:
        st.info("En modo manual no se guarda en la base de datos. "
                "Seleccione un equipo registrado para habilitar esta función.")
    else:
        st.info("El resultado no tiene un intervalo calculado, no hay nada que guardar.")

    # ── Exportar reporte ──────────────────────────────────────────────────────
    st.divider()
    st.subheader("6 · Exportar reporte")

    if intervalo_r is not None and fecha_base:
        fecha_proxima = fecha_desde_intervalo(fecha_base, intervalo_r)
        texto_ic      = formatear_intervalo(intervalo_r)

        reporte_docx = generar_reporte_ic_docx(
            equipo=equipo,
            unidad=unidad,
            emp=emp,
            met_usado=met_usado,
            intervalo_r=intervalo_r,
            texto_ic=texto_ic,
            fecha_proxima=fecha_proxima,
            recomendacion=recomendacion_guardada,
            msg_acotado=msg_acotado,
            resultados_puntos=resultados_puntos,
        )
        st.download_button(
            label     = "📄 Descargar reporte Word (.docx)",
            data      = reporte_docx,
            file_name = f"IC_{equipo['id'] if equipo else 'manual'}_{date.today()}.docx",
            mime      = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    # ── Historial de cálculos anteriores ─────────────────────────────────────
    if equipo and equipo.get("id"):
        hist = get_historial_ic(equipo["id"])
        if not hist.empty:
            st.divider()
            st.subheader("7 · Historial de cálculos IC")
            hist_display = hist.copy()
            hist_display["IC detallado"] = hist_display["intervalo_anios"].apply(
                lambda x: formatear_intervalo(x) if x else "—"
            )
            hist_display["ajustes_consecutivos"] = hist_display["ajustes_consecutivos"].apply(
                lambda x: f"{'↑' if x > 0 else '↓'} {abs(x)}x" if x != 0 else "—"
            )
            st.dataframe(
                hist_display.rename(columns={
                    "fecha_calculo":        "Fecha cálculo",
                    "metodo":               "Método",
                    "intervalo_anios":      "IC (años)",
                    "IC detallado":         "IC detallado",
                    "fecha_proxima_cal":    "Próxima calibración",
                    "ajustes_consecutivos": "Tendencia"
                })[["Fecha cálculo", "Método", "IC (años)", "IC detallado",
                     "Próxima calibración", "Tendencia"]],
                use_container_width=True,
                hide_index=True,
            )


# ══════════════════════════════════════════════════════════════════════════════
# INTERFAZ PRINCIPAL
# ══════════════════════════════════════════════════════════════════════════════

def mostrar_intervalos():
    """
    Función principal. Orquesta toda la interfaz del módulo.

    Flujo:
      1. Verifica Plan Premium
      2. Inicializa tabla resultados_ic en BD
      3. Encabezado del módulo
      4. Sección 1: selección del equipo (BD o manual)
      5. Parámetros técnicos (unidad, EMP, intervalo fabricante)
      6. Sección 2: selección del método
      7. Sección 3: formulario según el método
         → Validación de datos antes de habilitar el cálculo
      8. Al calcular: ejecuta el método, guarda en session_state
      9. Sección 4: resultados, gráficas, guardar, exportar, historial
    """
    aplicar_estilos()
    init_tabla_resultados_ic()   # Crea la tabla si no existe

    # ── Verificación Premium ──────────────────────────────────────────────────
    if "premium_activo" not in st.session_state or not st.session_state.premium_activo:
        with st.container(border=True):
            st.warning("🔒 Este módulo requiere el Plan Premium.")
            st.caption("Ve a **Acceso a Plan Premium** en el menú lateral para activarlo.")
        st.stop()

    # ── Encabezado ────────────────────────────────────────────────────────────
    st.markdown("""
    <div style="background:linear-gradient(135deg,#063d7d 0%,#238d93 100%);
    border-radius:12px;padding:1.4rem 2rem;margin-bottom:1.5rem;color:white;">
        <h1 style="margin:0;font-size:1.6rem;color:white;">📐 Intervalos de Calibración</h1>
        <p style="margin:0.3rem 0 0;opacity:0.85;font-size:0.9rem;">
        Determinación y ajuste del intervalo según ILAC-G24 / OIML D10:2007</p>
    </div>
    """, unsafe_allow_html=True)

    # ── Sección 1: Equipo ─────────────────────────────────────────────────────
    st.subheader("1 · Equipo a evaluar")

    modo = st.radio(
        "Modo de ingreso:",
        ["Seleccionar equipo registrado", "Ingresar datos manualmente"],
        horizontal=True
    )

    equipo = None

    if modo == "Seleccionar equipo registrado":
        df_equipos = get_equipos()
        df_cal     = get_calibraciones()

        if df_equipos.empty:
            st.warning("No hay equipos registrados. Primero agrega equipos en Ficha de equipo.")
            st.stop()

        equipos_disponibles = []
        for _, row in df_equipos.iterrows():
            cals = df_cal[df_cal["equipo_id"] == row["id"]].sort_values("fecha")
            historial = []
            for _, cal in cals.iterrows():
                historial.append({
                    "fecha":                     cal["fecha"],
                    "fue_ajustado":              False,
                    "errores_por_punto":         [{"valor_nominal": 0.0,
                                                   "error": cal["error"],
                                                   "incertidumbre": cal["incertidumbre"]}],
                    "metodo_ic_usado":           cal["metodo"],
                    "intervalo_calculado_anios": None,
                    "fecha_proxima_calibracion": cal["proximo_vencimiento"],
                    "calculado_por":             cal["laboratorio"]
                })

            intervalo_fabricante = 1.0
            if not cals.empty:
                ultima_cal = cals.iloc[-1]
                try:
                    fecha_ultima = datetime.strptime(ultima_cal["fecha"], "%Y-%m-%d").date()
                    fecha_proxima = datetime.strptime(
                        ultima_cal["proximo_vencimiento"], "%Y-%m-%d"
                    ).date()
                    intervalo_fabricante = max(
                        0.25, calcular_intervalo_anios(fecha_ultima, fecha_proxima)
                    )
                except Exception:
                    intervalo_fabricante = 1.0

            equipos_disponibles.append({
                "id":                         row["id"],
                "nombre":                     row["nombre"],
                "tipo":                       row["tipo"],
                "unidad":                     row["unidad"],
                "tolerancia_emp":             row["tolerancia"],
                "intervalo_fabricante_anios": intervalo_fabricante,
                "area":                       row.get("ubicacion", ""),
                "responsable":                row.get("responsable", ""),
                "estado":                     row["estado"],
                "historial_calibraciones":    historial
            })

        opciones  = {f"[{e['id']}] {e['nombre']}": e for e in equipos_disponibles}
        seleccion = st.selectbox("Seleccione el equipo:", list(opciones.keys()))
        equipo    = opciones[seleccion]

        st.markdown(f"""
        <div style="background:#0a453c33;border-left:4px solid #23c057;
        border-radius:8px;padding:0.9rem 1.2rem;margin-bottom:1rem;">
            <strong>{equipo['nombre']}</strong> &nbsp;|&nbsp;
            Tipo: {equipo.get('tipo','—')} &nbsp;|&nbsp;
            Área: {equipo.get('area','—')} &nbsp;|&nbsp;
            Estado: {equipo.get('estado','—')}
        </div>
        """, unsafe_allow_html=True)

        # ── Resumen del último cálculo IC guardado ────────────────────────────
        # Muestra contexto inmediato antes de recalcular. Si el equipo no tiene
        # ningún cálculo previo muestra un mensaje orientativo.
        hist_previo = get_historial_ic(equipo["id"])
        if not hist_previo.empty:
            ultimo      = hist_previo.iloc[0]
            ic_prev     = ultimo["intervalo_anios"]
            fecha_prox  = ultimo["fecha_proxima_cal"]
            metodo_prev = ultimo["metodo"]
            fecha_calc  = ultimo["fecha_calculo"]

            # Determinar si la próxima calibración ya venció o está próxima
            hoy = date.today()
            try:
                dias_restantes = (
                    datetime.strptime(fecha_prox, "%Y-%m-%d").date() - hoy
                ).days
                if dias_restantes < 0:
                    estado_fecha = f"🚨 Venció hace {abs(dias_restantes)} días"
                    color_fecha  = "#f4433622"
                    borde_fecha  = "#f44336"
                elif dias_restantes <= 30:
                    estado_fecha = f"⚠️ Vence en {dias_restantes} días"
                    color_fecha  = "#ff980022"
                    borde_fecha  = "#ff9800"
                else:
                    estado_fecha = f"✅ Vence en {dias_restantes} días"
                    color_fecha  = "#23c05722"
                    borde_fecha  = "#23c057"
                fecha_prox_fmt = datetime.strptime(fecha_prox, "%Y-%m-%d").strftime("%d/%m/%Y")
            except Exception:
                estado_fecha   = "—"
                color_fecha    = "#23c05722"
                borde_fecha    = "#23c057"
                fecha_prox_fmt = fecha_prox or "—"

            ic_detalle = formatear_intervalo(ic_prev) if ic_prev else "—"
            st.markdown(f"""
            <div style="background:{color_fecha};border-left:4px solid {borde_fecha};
            border-radius:8px;padding:0.8rem 1.2rem;margin-bottom:0.5rem;">
                <strong>📋 Último cálculo IC guardado</strong><br>
                Calculado el <strong>{fecha_calc}</strong> &nbsp;|&nbsp;
                Método: <strong>{metodo_prev}</strong><br>
                Intervalo: <strong>{ic_prev:.2f} años ({ic_detalle})</strong>
                &nbsp;|&nbsp;
                Próxima calibración: <strong>{fecha_prox_fmt}</strong><br>
                {estado_fecha}
            </div>
            """, unsafe_allow_html=True)
        else:
            st.info("ℹ️ Este equipo no tiene cálculos de IC guardados. "
                    "Complete el formulario a continuación para generar el primero.")

    contexto_actual = equipo["id"] if equipo and equipo.get("id") else "manual"
    if st.session_state.get("ic_contexto") != contexto_actual:
        for key in ("ic_resultado", "ic_fecha_base", "ic_metodo", "ic_fig", "ic_puntos", "ic_guardado_ok"):
            st.session_state.pop(key, None)
        st.session_state["ic_contexto"] = contexto_actual

    # ── Parámetros técnicos ───────────────────────────────────────────────────
    st.markdown("**Parámetros del instrumento:**")
    col1, col2, col3 = st.columns(3)

    with col1:
        unidad = st.text_input("Unidad de medida",
                               value=equipo.get("unidad", "") if equipo else "",
                               placeholder="mm, kg, °C…")
    with col2:
        emp_default = equipo.get("tolerancia_emp") if equipo else None
        emp = st.number_input(
            f"Tolerancia / {EMP_LABEL} ({unidad or 'unidad'})",
            min_value=0.0,
            value=float(emp_default) if emp_default else 0.02,
            format="%.5f", step=0.001
        )
        if equipo and emp_default not in (None, ""):
            st.caption(f"Valor precargado desde la tolerancia / {EMP_LABEL} registrado en la ficha del equipo.")
    with col3:
        fab_default = equipo.get("intervalo_fabricante_anios") if equipo else None
        intervalo_fabricante = st.number_input(
            "Intervalo fabricante (años)",
            min_value=0.1,
            value=float(fab_default) if fab_default else 1.0,
            format="%.2f", step=0.25
        )

    if mostrar_errores_validacion(validar_unidad("Unidad de medida", unidad)):
        st.stop()

    st.divider()

    # ── Sección 2: Método ─────────────────────────────────────────────────────
    st.subheader("2 · Método de evaluación")

    METODOS = {
        "Escalera — Error medio":
            f"Compara el **error medio** con el 80% del {EMP_LABEL}. "
            "Amplía o reduce el IC en 50% según el resultado. "
            "No aplica si el equipo fue ajustado recientemente.",
        "Escalera — Error con incertidumbre":
            f"Compara **|error| + incertidumbre expandida** con el 80% del {EMP_LABEL}. "
            "Más conservador que el error medio solo.",
        "Cartas de control":
            "Analiza la **deriva del error** en múltiples calibraciones. "
            f"Usa el **100% del {EMP_LABEL}** como límite. "
            "Requiere historial de al menos 2 fechas."
    }

    metodo = st.selectbox("Seleccione el método:", list(METODOS.keys()))
    st.info(METODOS[metodo])
    st.divider()

    # ── Sección 3: Formulario ─────────────────────────────────────────────────
    st.subheader("3 · Datos de calibración")

    resultado_calculo  = None
    fecha_cal_actual   = None
    resultados_puntos  = None

    # ── MÉTODOS 1 Y 2 ─────────────────────────────────────────────────────────
    if metodo in ("Escalera — Error medio", "Escalera — Error con incertidumbre"):

        fue_ajustado = st.selectbox(
            "¿El equipo fue ajustado mecánicamente en la última calibración?",
            ["No", "Sí"]
        ) == "Sí"

        col_a, col_b = st.columns(2)
        with col_a:
            fecha_cal_ant = st.date_input(
                "Fecha calibración anterior",
                value=date.today().replace(year=date.today().year - 1),
                min_value=FECHA_MIN,
                max_value=FECHA_MAX,
            )
        with col_b:
            fecha_cal_actual = st.date_input(
                "Fecha calibración actual (vigente)",
                value=date.today(),
                min_value=FECHA_MIN,
                max_value=FECHA_MAX,
            )

        intervalo_anterior = calcular_intervalo_anios(fecha_cal_ant, fecha_cal_actual)
        st.caption(
            f"Intervalo anterior: **{intervalo_anterior:.3f} años** "
            f"({formatear_intervalo(intervalo_anterior)})"
        )

        error_medio   = st.number_input(f"Error medio ({unidad})",
                                         value=0.01, format="%.5f", step=0.001)
        incertidumbre = 0.0
        if metodo == "Escalera — Error con incertidumbre":
            incertidumbre = st.number_input(f"Incertidumbre expandida U ({unidad})",
                                             min_value=0.0, value=0.002,
                                             format="%.5f", step=0.001)

        # Validación antes de mostrar el botón
        errores_validacion = validar_datos_escalera(
            error_medio, emp, incertidumbre, fecha_cal_ant, fecha_cal_actual
        )
        if errores_validacion:
            for msg in errores_validacion:
                st.error(f"❌ {msg}")
        else:
            if st.button("Calcular intervalo", type="primary"):
                st.session_state.pop("ic_guardado_ok", None)
                if metodo == "Escalera — Error medio":
                    resultado_calculo = escalera_error_medio(
                        error_medio, emp, intervalo_anterior, fue_ajustado
                    )
                    fig = grafica_escalera(error_medio, emp, unidad or "unidad")
                else:
                    resultado_calculo = escalera_error_incertidumbre(
                        error_medio, incertidumbre, emp, intervalo_anterior, fue_ajustado
                    )
                    fig = grafica_escalera(error_medio, emp, unidad or "unidad",
                                           con_incertidumbre=True, incertidumbre=incertidumbre)

                st.session_state["ic_resultado"]  = resultado_calculo
                st.session_state["ic_fecha_base"] = fecha_cal_actual
                st.session_state["ic_metodo"]     = metodo
                st.session_state["ic_fig"]        = fig
                st.session_state["ic_puntos"]     = None

    # ── MÉTODO 3 ──────────────────────────────────────────────────────────────
    else:
        fue_ajustado_cc = st.selectbox(
            "¿El equipo fue ajustado en alguna de las calibraciones?",
            ["No", "Sí"]
        ) == "Sí"

        fecha_ajuste_cc = None
        if fue_ajustado_cc:
            fecha_ajuste_cc = st.date_input(
                "Fecha del ajuste mecánico",
                min_value=FECHA_MIN,
                max_value=FECHA_MAX,
            )

        n_puntos = int(st.number_input("Cantidad de puntos críticos",
                                        min_value=1, max_value=10, value=1, step=1))

        historial_bd = []
        if equipo and equipo.get("id"):
            df_hist = get_calibraciones(equipo["id"])
            if not df_hist.empty:
                df_hist = df_hist.dropna(subset=["fecha", "error"]).sort_values("fecha")
                for _, cal in df_hist.iterrows():
                    try:
                        historial_bd.append({
                            "fecha": datetime.strptime(str(cal["fecha"]), "%Y-%m-%d").date(),
                            "error": float(cal["error"]),
                            "incertidumbre": float(cal["incertidumbre"]) if pd.notna(cal["incertidumbre"]) else 0.0,
                            "certificado": cal.get("numero_certificado", ""),
                        })
                    except Exception:
                        pass

        tiene_historial_suficiente = len(historial_bd) >= 2
        opciones_origen_hist = ["Ingreso manual"]
        if tiene_historial_suficiente:
            opciones_origen_hist.insert(0, "Usar historial registrado en Control de calibraciones")

        origen_historial = st.radio(
            "Origen del historial para cartas de control:",
            opciones_origen_hist,
            horizontal=True
        )

        fechas_cc = []
        errores_precargados = []

        if origen_historial == "Usar historial registrado en Control de calibraciones":
            if n_puntos > 1:
                st.info(
                    "Control de calibraciones registra un solo error por calibración. "
                    "Para varios puntos críticos, use ingreso manual."
                )
                n_puntos = 1
            st.success(
                f"Se usarán {len(historial_bd)} calibraciones registradas para este equipo."
            )
            st.dataframe(
                pd.DataFrame(historial_bd).rename(columns={
                    "fecha": "Fecha",
                    "error": "Error",
                    "incertidumbre": "Incertidumbre",
                    "certificado": "Certificado",
                }),
                use_container_width=True,
                hide_index=True,
            )
            fechas_cc = [item["fecha"] for item in historial_bd]
            errores_precargados = [item["error"] for item in historial_bd]
        else:
            if equipo and equipo.get("id") and not tiene_historial_suficiente:
                st.warning(
                    "Este equipo no tiene al menos 2 calibraciones registradas con fecha y error. "
                    "Ingrese el historial manualmente para aplicar cartas de control."
                )

            n_fechas = int(st.number_input("Cantidad de calibraciones en el historial",
                                            min_value=2, max_value=20, value=3, step=1))

            st.markdown("**Fechas de calibración:**")
            cols_fechas = st.columns(min(n_fechas, 4))
            for i in range(n_fechas):
                with cols_fechas[i % 4]:
                    default_f = date.today().replace(year=date.today().year - n_fechas + i)
                    f = st.date_input(
                        f"Fecha {i+1}",
                        value=default_f,
                        min_value=FECHA_MIN,
                        max_value=FECHA_MAX,
                        key=f"cc_fecha_{i}",
                    )
                    fechas_cc.append(f)

        fecha_cal_actual  = fechas_cc[-1]
        errores_validacion = validar_datos_cartas(fechas_cc, emp)

        if errores_validacion:
            for msg in errores_validacion:
                st.error(f"❌ {msg}")
        else:
            st.markdown("**Errores por punto crítico:**")
            errores_por_punto = []

            for p in range(n_puntos):
                with st.expander(f"Punto {p+1}", expanded=(p == 0)):
                    val_nominal = st.number_input(
                        f"Valor nominal P{p+1} ({unidad})",
                        min_value=0.0001, value=10.0,
                        format="%.4f", key=f"vn_{p}"
                    )
                    errores_p = []
                    for i, f in enumerate(fechas_cc):
                        es_ajuste = fue_ajustado_cc and fecha_ajuste_cc and f == fecha_ajuste_cc
                        if es_ajuste:
                            st.caption(f"  Fecha {i+1} ({f}) — Ajuste: error = 0 (automático)")
                            errores_p.append(0.0)
                        elif origen_historial == "Usar historial registrado en Control de calibraciones":
                            err = errores_precargados[i]
                            st.number_input(
                                f"Error en {f} ({unidad})",
                                value=float(err), format="%.5f",
                                step=0.001, key=f"err_hist_{p}_{i}",
                                disabled=True
                            )
                            errores_p.append(float(err))
                        else:
                            err = st.number_input(
                                f"Error en {f} ({unidad})",
                                value=0.0, format="%.5f",
                                step=0.001, key=f"err_{p}_{i}"
                            )
                            errores_p.append(err)
                    errores_por_punto.append({"nominal": val_nominal, "errores": errores_p})

            if st.button("Calcular intervalo", type="primary"):
                st.session_state.pop("ic_guardado_ok", None)
                resultados_puntos = []
                for p_data in errores_por_punto:
                    res = cartas_control(
                        fechas_cc, p_data["errores"], emp,
                        fue_ajustado_cc, fecha_ajuste_cc
                    )
                    resultados_puntos.append({"nominal": p_data["nominal"], **res})

                intervalos_validos = [
                    r["intervalo_nuevo"] for r in resultados_puntos
                    if r["intervalo_nuevo"] is not None
                ]

                if intervalos_validos:
                    intervalo_final = min(intervalos_validos)
                    tipo_final      = "ok"
                else:
                    intervalo_final = None
                    tipo_final      = "bad" if any(r["tipo"] == "bad"
                                                   for r in resultados_puntos) else "warn"

                texto_final = formatear_intervalo(intervalo_final) if intervalo_final else "—"
                resultado_calculo = {
                    "recomendacion": (
                        f"Intervalo más restrictivo entre {n_puntos} punto(s): "
                        f"{intervalo_final:.2f} años ({texto_final})."
                        if intervalo_final else
                        "No se pudo calcular un intervalo válido. Revise los datos."
                    ),
                    "intervalo_nuevo": intervalo_final,
                    "tipo":           tipo_final
                }

                st.session_state["ic_resultado"]  = resultado_calculo
                st.session_state["ic_fecha_base"] = fecha_cal_actual
                st.session_state["ic_metodo"]     = metodo
                st.session_state["ic_fig"]        = None
                st.session_state["ic_puntos"]     = resultados_puntos

    # ── Sección 4: Resultados ─────────────────────────────────────────────────
    if "ic_resultado" in st.session_state and st.session_state["ic_resultado"]:
        st.divider()
        st.subheader("4 · Resultados")
        mostrar_resultado(
            res                  = st.session_state["ic_resultado"],
            fecha_base           = st.session_state.get("ic_fecha_base"),
            met_usado            = st.session_state.get("ic_metodo", metodo),
            equipo               = equipo,
            unidad               = unidad,
            emp                  = emp,
            intervalo_fabricante = intervalo_fabricante,
            resultados_puntos    = st.session_state.get("ic_puntos")
        )


# ── Punto de entrada ──────────────────────────────────────────────────────────
mostrar_intervalos()
